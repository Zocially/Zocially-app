import os
import google.generativeai as genai
import PyPDF2
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import google.api_core.exceptions

class CVProcessor:
    def __init__(self):
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError("GOOGLE_API_KEY not found in environment variables")
        
        genai.configure(api_key=api_key)
        
        # Configure safety settings to prevent blocking professional CV content
        from google.generativeai.types import HarmCategory, HarmBlockThreshold
        
        safety_settings = [
            {
                "category": HarmCategory.HARM_CATEGORY_HARASSMENT,
                "threshold": HarmBlockThreshold.BLOCK_NONE,
            },
            {
                "category": HarmCategory.HARM_CATEGORY_HATE_SPEECH,
                "threshold": HarmBlockThreshold.BLOCK_NONE,
            },
            {
                "category": HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
                "threshold": HarmBlockThreshold.BLOCK_NONE,
            },
            {
                "category": HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
                "threshold": HarmBlockThreshold.BLOCK_NONE,
            },
        ]
        
        # Switch to 2.0-flash-exp (available and free tier)
        # Switch to gemini-flash-latest (Explicitly available in user list)
        self.model = genai.GenerativeModel(
            'gemini-flash-latest',
            safety_settings=safety_settings
        )

    def extract_text(self, file_path):
        """Extracts text from PDF."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
            return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""

    @retry(
        retry=retry_if_exception_type(google.api_core.exceptions.ResourceExhausted),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    def assess_cv(self, cv_text):
        """Assess the CV against general best practices."""
        prompt = f"""
        Act as an expert career coach. Review the following CV and provide a brief assessment.
        Highlight 3 strengths and 3 areas for improvement.
        
        CV Content:
        {cv_text}
        """
        response = self.model.generate_content(prompt)
        return response.text
    
    def identify_cv_gaps(self, cv_text):
        """
        Identifies missing or weak elements in the CV.
        Returns a dict with gaps and user-friendly prompts.
        """
        import re
        
        gaps = {
            'has_gaps': False,
            'missing_elements': [],
            'prompts': {}
        }
        
        # Check for phone number
        has_phone = bool(re.search(r'\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}', cv_text))
        if not has_phone:
            gaps['missing_elements'].append('phone')
            gaps['prompts']['phone'] = "📞 Phone Number (e.g., (555) 123-4567)"
            gaps['has_gaps'] = True
        
        # Check for LinkedIn
        has_linkedin = bool(re.search(r'linkedin\.com/in/[\w-]+', cv_text, re.IGNORECASE))
        if not has_linkedin:
            gaps['missing_elements'].append('linkedin')
            gaps['prompts']['linkedin'] = "🔗 LinkedIn Profile URL (e.g., https://linkedin.com/in/yourname)"
            gaps['has_gaps'] = True
        
        # Check for location
        has_location = bool(re.search(r'\b[A-Z][a-z]+,\s*[A-Z]{2}\b', cv_text))
        if not has_location:
            gaps['missing_elements'].append('location')
            gaps['prompts']['location'] = "📍 Location (e.g., New York, NY or London, UK)"
            gaps['has_gaps'] = True
        
        # Check for professional summary
        has_summary = bool(re.search(r'##\s*(Professional\s+)?Summary', cv_text, re.IGNORECASE))
        summary_text = ""
        if has_summary:
            summary_match = re.search(r'##\s*(Professional\s+)?Summary\s*\n(.+?)(?=\n##|\Z)', cv_text, re.IGNORECASE | re.DOTALL)
            if summary_match:
                summary_text = summary_match.group(2).strip()
        
        if not has_summary or len(summary_text) < 50:
            gaps['missing_elements'].append('summary')
            gaps['prompts']['summary'] = "💼 Professional Summary (Brief overview of your career and goals)"
            gaps['has_gaps'] = True
        
        # Check for skills section
        has_skills = bool(re.search(r'##\s*(Skills|Core\s+Competencies|Technical\s+Skills)', cv_text, re.IGNORECASE))
        if not has_skills:
            gaps['missing_elements'].append('skills')
            gaps['prompts']['skills'] = "🎯 Key Skills (Comma-separated, e.g., Python, Project Management, Communication)"
            gaps['has_gaps'] = True
        
        # Check for Work Experience
        has_work_exp = bool(re.search(r'##\s*(Work\s+Experience|Professional\s+Experience)', cv_text, re.IGNORECASE))
        if not has_work_exp:
             gaps['missing_elements'].append('work_experience')
             gaps['prompts']['work_experience'] = "💼 Work Experience (e.g., Job Title | Company | Dates | Key Responsibilities)"
             gaps['has_gaps'] = True

        # Check for Education
        has_education = bool(re.search(r'##\s*Education', cv_text, re.IGNORECASE))
        if not has_education:
             gaps['missing_elements'].append('education')
             gaps['prompts']['education'] = "🎓 Education (e.g., Degree | University | Year)"
             gaps['has_gaps'] = True

        # Check for quantifiable achievements
        numbers_count = len(re.findall(r'\d+%|\$\d+|\d+\+', cv_text))
        if numbers_count < 3:
            gaps['missing_elements'].append('achievements')
            gaps['prompts']['achievements'] = "📊 Key Achievements (Include numbers/metrics, e.g., 'Increased sales by 30%')"
            gaps['has_gaps'] = True
        
        return gaps

    @retry(
        retry=retry_if_exception_type(google.api_core.exceptions.ResourceExhausted),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    def tailor_cv(self, cv_text, job_description, additional_info=None):
        """Rewrites the CV to match the job description with ATS optimization."""
        
        # Format additional information if provided
        additional_info_text = ""
        if additional_info:
            additional_info_text = "\\n\\n**ADDITIONAL INFORMATION PROVIDED BY USER:**\\n"
            additional_info_text += "The user has provided the following additional information to enhance their CV:\\n"
            if additional_info.get('phone'):
                additional_info_text += f"- Phone: {additional_info['phone']}\\n"
            if additional_info.get('linkedin'):
                additional_info_text += f"- LinkedIn: {additional_info['linkedin']}\\n"
            if additional_info.get('location'):
                additional_info_text += f"- Location: {additional_info['location']}\\n"
            if additional_info.get('summary'):
                additional_info_text += f"- Professional Summary: {additional_info['summary']}\\n"
            if additional_info.get('skills'):
                additional_info_text += f"- Additional Skills: {additional_info['skills']}\\n"
            if additional_info.get('achievements'):
                additional_info_text += f"- Key Achievements: {additional_info['achievements']}\\n"
            additional_info_text += "\\nPlease incorporate this information naturally into the tailored CV."
        
        prompt = f"""
        Act as an expert CV writer specializing in ATS (Applicant Tracking System) optimization. Tailor the following CV to match the Job Description provided.
        
        **CRITICAL INSTRUCTION ON DATES:**
        You must PRESERVE all dates exactly as they appear in the original CV. 
        - Do NOT change "Jan 2020" to "January 2020".
        - Do NOT change "2020 - Present" to "2020 - 2023".
        - Keep the exact date strings for every work experience and education entry.

        **CRITICAL RULES ON FACTUAL ACCURACY (ZERO HALLUCINATION):**
        1. **DO NOT INVENT INFORMATION:** You must NOT add any Education, Work Experience, Job Titles, or Companies that are not explicitly present in the original CV.
        2. **MISSING SECTIONS:** If the original CV does not have an Education section, DO NOT CREATE ONE. It is better to have a missing section than a fake one.
        3. **ONLY TAILOR EXISTING CONTENT:** You can rephrase responsibilities to match keywords, but you cannot invent new responsibilities or skills that the candidate clearly does not possess based on the text.

        **ATS OPTIMIZATION REQUIREMENTS:**
        1. **Contact Information (Header Section):**
           - **Phone:** Use standard format: (XXX) XXX-XXXX or XXX-XXX-XXXX
           - **Email:** Professional email address on its own line or clearly separated
           - **LinkedIn:** Use full URL (https://linkedin.com/in/username), not shortened links
           - **Location:** Format as "City, State" or "City, Country" (e.g., "London, UK" or "New York, NY")
           - Ensure all contact info is on separate lines or separated by " | "
        
        2. **Keyword Matching:**
           - Extract key skills, technologies, and qualifications from the job description
           - Naturally incorporate these keywords throughout the CV (especially in Skills and Work Experience)
           - Match exact terminology used in the job posting (e.g., if they say "JavaScript" don't say "JS")
           - Include both acronyms and full terms on FIRST mention (e.g., "Artificial Intelligence (AI)")
           - Subsequent mentions can use acronym only
        
        3. **Standard Section Headers (ATS-Recognizable):**
           Use ONLY these exact section headers:
           - ## Professional Summary (or ## Summary)
           - ## Core Competencies (or ## Skills or ## Technical Skills)
           - ## Work Experience (or ## Professional Experience)
           - ## Education (ONLY IF present in original CV)
           - ## Certifications (if applicable)
           - ## Projects (if applicable)
           DO NOT use creative headers like "Career Journey" or "My Expertise"
        
        4. **Quantifiable Achievements:**
           - Include numbers, percentages, and metrics wherever possible
           - Use action verbs (Led, Developed, Increased, Reduced, Managed, Implemented, Achieved, etc.)
           - Format: "Action Verb + Task + Quantifiable Result"
           - Example: "Increased sales by 35% through implementation of new CRM system"
        
        5. **Skills Section (CRITICAL FOR ATS):**
           - Create a dedicated "Core Competencies" or "Technical Skills" section
           - List skills in order of relevance to the job description (most important first)
           - Group related skills with clear labels:
             * **Programming Languages:** Python, JavaScript, Java
             * **Frameworks & Tools:** React, Node.js, Docker
             * **Soft Skills:** Leadership, Communication, Problem-solving
           - Include skill variations where relevant (e.g., "JavaScript (JS, ES6+)")
           - Use comma-separated format or bullet points (NO tables)
           - Include both hard skills (technical) and soft skills mentioned in job posting
        
        6. **Date Format Consistency:**
           - Use consistent format throughout: "Month YYYY - Month YYYY" (e.g., "January 2020 - March 2023")
           - Alternative acceptable format: "MM/YYYY - MM/YYYY"
           - Use "Present" for current positions (not "Current" or "Now")
           - ALWAYS include date ranges for ALL work experience entries
           - Ensure dates are in reverse chronological order (most recent first)

        **FORMATTING GUIDELINES (ATS-Friendly):**
        1. **Structure:** Follow this standard professional format:
           - **Header:** Name, Contact Info (Email, Phone, LinkedIn, Location).
           - **Professional Summary:** A strong, tailored summary (3-4 lines) aligning with the job.
           - **Skills:** A concise list of relevant hard and soft skills.
           - **Work Experience:** Reverse chronological order. For each role include **Job Title**, **Company**, **Location**, and **Date Range**.
           - **Education:** Degree, University, Year.
           - **Projects/Certifications:** (If applicable and relevant).
        
        2. **Content:** 
           - Use simple, clear language (avoid jargon unless it's in the job description)
           - Keep bullet points concise (1-2 lines each)
           - Focus on achievements and impact, not just responsibilities
        
        3. **Tone:** Professional, confident, and action-oriented.
        
        4. **Format:** Return ONLY the content of the new CV, formatted in clean Markdown. 
           - Use `##` for section headers (e.g., ## Professional Summary).
           - Use `###` for sub-headers (e.g., ### Software Engineer | Google).
           - Use `**bold**` for key terms (job titles, company names, key achievements).
           - Use `*italic*` sparingly for less emphatic emphasis if needed.
           - Use `-` for bullet points.
           - DO NOT use tables, text boxes, or complex formatting.
           - DO NOT include images, graphics, or special characters.

        **SPELLING AND GRAMMAR (CRITICAL):**
        - **Proofread thoroughly:** Fix ALL spelling, grammar, and punctuation errors
        - **Common mistakes to avoid:**
          * Their/There/They're, Your/You're, Its/It's
          * Affect/Effect, Then/Than
          * Ensure/Insure, Complement/Compliment
        - **Grammar rules:**
          * Use consistent verb tense (past tense for previous roles, present for current)
          * Ensure subject-verb agreement
          * Avoid run-on sentences
          * Use proper punctuation (commas, periods, semicolons)
        - **Professional language:**
          * No slang or informal language
          * No contractions (use "do not" instead of "don't")
          * Capitalize proper nouns (company names, job titles, locations)
          * Use active voice, not passive
        - **Consistency:**
          * Consistent date formats throughout
          * Consistent bullet point style
          * Consistent capitalization in section headers

        **ATS COMPATIBILITY CHECKLIST:**
        - ✓ Use standard fonts (will be converted to Calibri in DOCX)
        - ✓ Use standard section headers
        - ✓ Include relevant keywords from job description
        - ✓ Use simple formatting (no tables, columns, or text boxes)
        - ✓ Include quantifiable achievements
        - ✓ Use chronological format
        - ✓ Spell out acronyms on first use

        CV Content:
        {cv_text}
        
        Job Description:
        {job_description}
        {additional_info_text}
        """
        
        try:
            response = self.model.generate_content(prompt)
            
            # Check if response was blocked
            if not response.parts:
                # Check finish_reason
                if hasattr(response, 'candidates') and response.candidates:
                    finish_reason = response.candidates[0].finish_reason
                    if finish_reason == 1:  # SAFETY
                        raise ValueError(
                            "The AI safety filters blocked the response. This can happen with very long CVs or job descriptions. "
                            "Try shortening your CV or job description, or try again in a moment."
                        )
                    elif finish_reason == 3:  # RECITATION
                        raise ValueError(
                            "The response was blocked due to potential copyright issues. "
                            "Please ensure your CV and job description don't contain copyrighted material."
                        )
                    else:
                        raise ValueError(
                            f"The AI model couldn't generate a response (finish_reason: {finish_reason}). "
                            "Please try again or contact support."
                        )
                else:
                    raise ValueError(
                        "The AI model returned an empty response. This might be due to safety filters or API issues. "
                        "Please try again in a moment."
                    )
            
            return response.text
        except Exception as e:
            # Re-raise with more context if it's not already our custom error
            if "safety filters" in str(e) or "finish_reason" in str(e):
                raise
            else:
                raise ValueError(f"Error generating tailored CV: {str(e)}")
    
    @retry(
        retry=retry_if_exception_type(google.api_core.exceptions.ResourceExhausted),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    def improve_cv_for_ats(self, cv_text, validation_report):
        """
        Improves a CV to achieve 90+ ATS score based on validation recommendations.
        """
        recommendations_text = "\n".join(validation_report.get('recommendations', []))
        current_score = validation_report.get('score', 0)
        
        prompt = f"""
        Act as an expert ATS optimization specialist. You have a CV that scored {current_score}/100 on ATS compatibility.
        Your goal is to improve this CV to achieve a score of 90+ (Grade A or B) while preserving all factual information.
        
        **CURRENT ISSUES TO FIX:**
        {recommendations_text}
        
        **CRITICAL ATS OPTIMIZATION REQUIREMENTS:**
        1. **Contact Information (MUST HAVE):**
           - Add standard phone format if missing: (XXX) XXX-XXXX
           - Add full LinkedIn URL if missing: https://linkedin.com/in/username
           - Add location in format: City, State (e.g., "New York, NY")
        
        2. **Section Headers (MUST USE EXACT NAMES):**
           - ## Professional Summary (or ## Summary)
           - ## Core Competencies (or ## Skills)
           - ## Professional Experience (or ## Work Experience)
           - ## Education
           - ## Certifications (if applicable)
        
        3. **Professional Summary (MUST HAVE 3-4 LINES):**
           - Strong opening statement
           - Key qualifications and years of experience
           - Core expertise areas
           - Career objective or value proposition
        
        4. **Skills Section (CRITICAL):**
           - Group skills by category:
             * **Technical Skills:** List 5-10 relevant technical skills
             * **Professional Skills:** List 3-5 soft skills
           - Use comma-separated format
           - Include industry-standard terminology
        
        5. **Quantifiable Achievements (ADD NUMBERS):**
           - Every work experience bullet should have metrics when possible
           - Use percentages, dollar amounts, or quantities
           - Format: "Action Verb + Task + Quantifiable Result"
           - Examples: "Increased revenue by 25%", "Managed team of 10", "Reduced costs by $50K"
        
        6. **Formatting (STRICT RULES):**
           - Use ## for main section headers
           - Use ### for job titles/companies
           - Use - for bullet points (no indentation)
           - Use **bold** for job titles and company names
           - Keep lines under 120 characters
           - Use Title Case for all section headers
        
        7. **Action Verbs (START EVERY BULLET):**
           Use strong action verbs: Led, Developed, Implemented, Achieved, Managed, Created, 
           Designed, Improved, Increased, Reduced, Streamlined, Optimized, etc.
        
        **INSTRUCTIONS:**
        - Fix ALL issues mentioned in the recommendations
        - Ensure EVERY required section is present
        - Add missing contact information (use placeholder if needed)
        - Add quantifiable metrics to work experience
        - Use standard ATS-friendly formatting
        - Maintain all factual information from original CV
        - Return ONLY the improved CV in Markdown format
        
        **ORIGINAL CV:**
        {cv_text}
        """
        
        response = self.model.generate_content(prompt)
        return response.text

    def validate_cv(self, cv_text: str) -> str:
        """Simple validation of the generated CV.
        Checks for presence of key sections and returns a report.
        """
        import re
        required_sections = [
            # "Header", # Header is usually implicit at the top
            "Professional Summary",
            "Skills",
            "Work Experience",
            "Education",
        ]
        missing = []
        for sec in required_sections:
            pattern = rf"##\s*{re.escape(sec)}"
            if not re.search(pattern, cv_text, re.IGNORECASE):
                missing.append(sec)
        if missing:
            return "Missing sections: " + ", ".join(missing)
        return "CV looks good – all required sections are present."
    
    def validate_ats_compatibility(self, cv_text, job_description=""):
        """
        Validates CV for ATS compatibility and returns a detailed report.
        Returns a dict with 'score', 'passed', and 'recommendations'.
        """
        import re
        
        score = 0
        max_score = 100
        recommendations = []
        
        # 1. Check for required sections (20 points)
        required_sections = ["Professional Summary", "Summary", "Skills", "Core Competencies", 
                           "Work Experience", "Professional Experience", "Education"]
        found_sections = []
        for sec in required_sections:
            if re.search(rf"##\s*{re.escape(sec)}", cv_text, re.IGNORECASE):
                found_sections.append(sec)
        
        if any(s in found_sections for s in ["Professional Summary", "Summary"]):
            score += 5
        else:
            recommendations.append("❌ Missing 'Professional Summary' section")
        
        if any(s in found_sections for s in ["Skills", "Core Competencies"]):
            score += 5
        else:
            recommendations.append("❌ Missing 'Skills' or 'Core Competencies' section")
        
        if any(s in found_sections for s in ["Work Experience", "Professional Experience"]):
            score += 5
        else:
            recommendations.append("❌ Missing 'Work Experience' section")
        
        if "Education" in found_sections:
            score += 5
        else:
            recommendations.append("❌ Missing 'Education' section")
        
        # 2. Check contact information format (20 points)
        has_email = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', cv_text))
        has_phone = bool(re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', cv_text))
        has_linkedin = bool(re.search(r'linkedin\.com/in/[\w-]+', cv_text, re.IGNORECASE))
        
        if has_email:
            score += 7
        else:
            recommendations.append("❌ No email address found")
        
        if has_phone:
            score += 7
        else:
            recommendations.append("⚠️ No phone number found or incorrect format")
        
        if has_linkedin:
            score += 6
        else:
            recommendations.append("⚠️ No LinkedIn profile found")
        
        # 3. Check for quantifiable achievements (15 points)
        numbers_count = len(re.findall(r'\d+%|\$\d+|\d+\+', cv_text))
        if numbers_count >= 5:
            score += 15
        elif numbers_count >= 3:
            score += 10
            recommendations.append("⚠️ Add more quantifiable achievements (numbers, percentages)")
        else:
            score += 5
            recommendations.append("❌ Very few quantifiable achievements found")
        
        # 4. Check for action verbs (10 points)
        action_verbs = ['Led', 'Developed', 'Managed', 'Implemented', 'Achieved', 'Increased', 
                       'Reduced', 'Created', 'Designed', 'Improved']
        verb_count = sum(1 for verb in action_verbs if verb in cv_text)
        if verb_count >= 5:
            score += 10
        elif verb_count >= 3:
            score += 7
        else:
            score += 3
            recommendations.append("⚠️ Use more action verbs (Led, Developed, Managed, etc.)")
        
        # 5. Check for keyword matching if job description provided (20 points)
        if job_description:
            # Extract potential keywords (simple approach)
            job_keywords = set(re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', job_description))
            job_keywords.update(re.findall(r'\b[A-Z]{2,}\b', job_description))  # Acronyms
            
            matched_keywords = [kw for kw in job_keywords if kw in cv_text]
            match_rate = len(matched_keywords) / max(len(job_keywords), 1)
            
            if match_rate >= 0.5:
                score += 20
            elif match_rate >= 0.3:
                score += 15
                recommendations.append("⚠️ Include more keywords from job description")
            else:
                score += 5
                recommendations.append("❌ Low keyword match with job description")
        else:
            score += 10  # Give partial credit if no job description provided
        
        # 6. Check file size (text length as proxy) (5 points)
        if len(cv_text) < 10000:  # Reasonable CV length
            score += 5
        else:
            recommendations.append("⚠️ CV might be too long (keep under 2 pages)")
        
        # 7. Check for problematic formatting (10 points)
        if '|' not in cv_text or cv_text.count('|') < 10:  # Likely no tables
            score += 5
        else:
            recommendations.append("⚠️ Possible table formatting detected (not ATS-friendly)")
        
        if not re.search(r'[^\x00-\x7F]', cv_text):  # No special characters
            score += 5
        else:
            recommendations.append("⚠️ Special characters detected (may cause ATS issues)")
        
        # 8. Check for proper markdown formatting (10 points)
        # Check for section headers (##)
        section_headers = re.findall(r'^##\s+.+$', cv_text, re.MULTILINE)
        if len(section_headers) >= 4:  # At least 4 main sections
            score += 5
        else:
            recommendations.append("❌ Missing proper section headers (use ## for sections)")
        
        # Check for consistent bullet points
        bullet_lines = re.findall(r'^\s*-\s+.+$', cv_text, re.MULTILINE)
        if bullet_lines:
            score += 5
        else:
            recommendations.append("⚠️ No bullet points found (use - for lists)")
        
        # 9. Check for indentation and formatting consistency (5 points)
        lines = cv_text.split('\n')
        inconsistent_indent = False
        for line in lines:
            if line.startswith('  -') or line.startswith('\t-'):  # Improper indentation
                inconsistent_indent = True
                break
        
        if not inconsistent_indent:
            score += 5
        else:
            recommendations.append("⚠️ Inconsistent indentation detected")
        
        # 10. Check for proper capitalization in headers (5 points)
        proper_caps = True
        for header in section_headers:
            # Remove ## and check if title case or all caps
            header_text = header.replace('##', '').strip()
            if not (header_text.istitle() or header_text.isupper()):
                proper_caps = False
                break
        
        if proper_caps and section_headers:
            score += 5
        else:
            recommendations.append("⚠️ Section headers should use Title Case")
        
        # 11. Check for reasonable line lengths (5 points)
        long_lines = [line for line in lines if len(line) > 120 and not line.startswith('http')]
        if len(long_lines) < 3:
            score += 5
        else:
            recommendations.append("⚠️ Some lines are too long (keep under 120 characters)")
        
        # Generate final report
        passed = score >= 70
        grade = "A" if score >= 90 else "B" if score >= 80 else "C" if score >= 70 else "D" if score >= 60 else "F"
        
        report = {
            "score": score,
            "grade": grade,
            "passed": passed,
            "recommendations": recommendations
        }
        
        return report

    def generate_docx(self, cv_text, filename):
        """Converts Markdown CV text to a professionally formatted DOCX file."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()
        
        # --- Styles ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Helper to add styled paragraph
        def add_styled_para(text, bold=False, italic=False, size=None, align=None, space_after=None):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            if align:
                p.alignment = align
            if space_after is not None:
                p.paragraph_format.space_after = Pt(space_after)
            return p

        # --- Parsing ---
        # Clean up markdown code blocks if present
        cv_text = cv_text.replace("```markdown", "").replace("```", "").strip()
        
        lines = cv_text.split('\n')
        
        # 1. Extract Header (Name & Contact)
        # Assumption: The first H1 is the Name, followed by contact lines until the next Header.
        header_lines = []
        body_lines = []
        
        is_header = True
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # If we hit a section header (##), header is done
            if line.startswith('## '):
                is_header = False
            
            if is_header:
                # Remove markdown headers for the name
                clean_line = line.replace('# ', '').replace('**', '').strip()
                header_lines.append(clean_line)
            else:
                body_lines.append(line)

        # 2. Write Header
        if header_lines:
            # Name (First line)
            name = header_lines[0]
            add_styled_para(name, bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
            
            # Contact Info (Joined by |)
            contact_info = " | ".join(header_lines[1:])
            add_styled_para(contact_info, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
            
            # Horizontal Line
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom_color = RGBColor(0, 0, 0)
            p.paragraph_format.border_bottom_width = Pt(1)
            p.paragraph_format.space_after = Pt(12)

        # 3. Write Body Sections
        current_section = ""
        skills_buffer = []
        
        def flush_skills():
            nonlocal skills_buffer
            if skills_buffer:
                # Join with a bullet point separator
                skill_text = " • ".join(skills_buffer)
                p = doc.add_paragraph(skill_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
                skills_buffer = []

        for line in body_lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if we are hitting a header or new section, flush skills if needed
            if line.startswith('## ') or line.startswith('### '):
                flush_skills()

            # Section Headers (##)
            if line.startswith('## '):
                current_section = line[3:].strip()
                p = doc.add_heading(current_section.upper(), level=1)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                continue
                
            # Sub-headers (###)
            if line.startswith('### '):
                text = line[4:].strip().replace('**', '')
                p = doc.add_heading(text, level=2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                continue
                
            # Bullet Points
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip().replace('**', '') # Simple bold removal for cleanliness
                
                # Special formatting for Skills section
                if "SKILLS" in current_section.upper():
                    skills_buffer.append(text)
                else:
                    # Normal bullet point
                    p = doc.add_paragraph(text, style='List Bullet')
            
            # Normal Text
            else:
                # If we encounter normal text in skills section, flush buffer first
                if "SKILLS" in current_section.upper() and skills_buffer:
                    flush_skills()

                # Handle bolding within text **like this**
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        # Final flush
        flush_skills()

        doc.save(filename)
        return filename


    @retry(
        retry=retry_if_exception_type(google.api_core.exceptions.ResourceExhausted),
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10)
    )
    def generate_cover_letter(self, cv_text, job_info):
        """Generates a cover letter using job info (title, company, description)."""
        # Support both dict and plain description string
        if isinstance(job_info, dict):
            title = job_info.get('title', '')
            company = job_info.get('company', '')
            description = job_info.get('description', '')
            summary = job_info.get('summary', '')
        else:
            title = ''
            company = ''
            description = job_info
            summary = ''
        # Truncate description to avoid overly long prompts
        max_len = 1500
        if len(description) > max_len:
            description = description[:max_len] + "..."
        # Build a detailed prompt, optionally including a summary
        prompt = f"""
        You are a professional cover letter writer.
        Write a concise, targeted cover letter for the position \"{title}\" at \"{company}\".
        Use the candidate's CV content and align with the job description below.
        """
        if summary:
            prompt += f"\nAdditional summary of the role (highlights, required skills, responsibilities):\n{summary}\n"
        prompt += f"""
        CV Content:
        {cv_text}
        
        Job Description:
        {description}
        """
        response = self.model.generate_content(prompt)
        return response.text
