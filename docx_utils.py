import io
from docx import Document
from docx.shared import Pt

def create_docx_from_markdown(markdown_text):
    """
    Converts simple Markdown text to a docx file in memory.
    Handles headers (#) and list items (- or *).
    """
    doc = Document()
    
    # Set default font (optional, but good for consistency)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in markdown_text.splitlines():
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Main Title (Name)
            p = doc.add_heading(line[2:], level=1)
            p.alignment = 1  # Center align
        elif line.startswith('## '):
            # Section Headers
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Sub-headers
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # List Items
            p = doc.add_paragraph(style='List Bullet')
            process_bold(p, line[2:])
        else:
            # Normal Paragraph
            p = doc.add_paragraph()
            process_bold(p, line)
            
    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def process_bold(paragraph, text):
    """
    Helper to parse **bold** and *italic* text in a string and add runs to a paragraph.
    Supports nested or mixed formatting like **bold** and *italic*.
    """
    # Simple parser: split by ** first, then by *
    # This is a basic implementation and might not handle complex nesting perfectly
    parts = text.split('**')
    for i, part in enumerate(parts):
        is_bold = (i % 2 == 1)
        
        # Now split by * for italics
        sub_parts = part.split('*')
        for j, sub_part in enumerate(sub_parts):
            is_italic = (j % 2 == 1)
            
            if sub_part:
                run = paragraph.add_run(sub_part)
                run.bold = is_bold
                run.italic = is_italic
